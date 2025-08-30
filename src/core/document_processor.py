# FIXED: document_processor.py - Enhanced Chunking Strategy

"""
Enhanced document processor with improved chunking for better RAG performance.
"""

import logging
import io
from typing import List, Optional, Tuple
from pypdf import PdfReader
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from ..config.settings import AIConfig, FileConfig
from .exceptions import DocumentProcessingError

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Enhanced document processor with semantic-aware chunking."""

    def __init__(self):
        self.ai_config = AIConfig()
        self.file_config = FileConfig()
        self.text_splitter = self._create_optimized_text_splitter()

    def _create_optimized_text_splitter(self) -> RecursiveCharacterTextSplitter:
        """Create optimized text splitter with better separators for semantic chunks."""
        return RecursiveCharacterTextSplitter(
            chunk_size=self.ai_config.CHUNK_SIZE,  # 800 from optimized settings
            chunk_overlap=self.ai_config.CHUNK_OVERLAP,  # 150 from optimized settings
            length_function=len,
            separators=[
                "\n\n\n",  # Major section breaks
                "\n\n",    # Paragraph breaks
                "\n",      # Line breaks
                ". ",      # Sentence endings (crucial for semantic integrity)
                "! ",      # Exclamation endings
                "? ",      # Question endings
                "; ",      # Semicolon breaks
                ", ",      # Comma breaks (last resort)
                " ",       # Word breaks (last resort)
                ""         # Character breaks (absolute last resort)
            ],
            keep_separator=True  # Keep separators to maintain context
        )

    def extract_text(self, file) -> Tuple[str, str]:
        """Extract text with enhanced cleaning."""
        if not file:
            raise DocumentProcessingError("No file provided")

        file_extension = file.name.split('.')[-1].lower()
        if file_extension not in self.file_config.ALLOWED_TYPES:
            raise DocumentProcessingError(
                f"Unsupported file type: {file_extension}. "
                f"Supported types: {', '.join(self.file_config.ALLOWED_TYPES)}"
            )

        # Check file size
        if file.size > self.file_config.MAX_FILE_SIZE * 1024 * 1024:
            raise DocumentProcessingError(
                f"File size ({file.size / 1024 / 1024:.1f}MB) exceeds "
                f"maximum allowed size ({self.file_config.MAX_FILE_SIZE}MB)"
            )

        try:
            if file_extension == 'pdf':
                text = self._extract_from_pdf(file)
            elif file_extension == 'docx':
                text = self._extract_from_docx(file)
            elif file_extension == 'txt':
                text = self._extract_from_txt(file)
            else:
                raise DocumentProcessingError(f"Unsupported file type: {file_extension}")

            if not text or not text.strip():
                raise DocumentProcessingError("No text content found in document")

            # Enhanced text cleaning
            text = self._enhanced_text_cleaning(text)

            # Check text length
            if len(text) > self.file_config.MAX_TEXT_LENGTH:
                logger.warning(f"Text length ({len(text)}) exceeds maximum, truncating")
                text = text[:self.file_config.MAX_TEXT_LENGTH]

            logger.info(f"Successfully extracted {len(text)} characters from {file.name}")
            return text.strip(), file_extension

        except Exception as e:
            logger.error(f"Text extraction failed for {file.name}: {e}")
            raise DocumentProcessingError(f"Failed to extract text from {file.name}: {str(e)}")

    def _enhanced_text_cleaning(self, text: str) -> str:
        """Enhanced text cleaning with better structure preservation."""
        import re
        
        # Remove excessive whitespace while preserving paragraph structure
        # First, normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive spaces but preserve single spaces
        text = re.sub(r' +', ' ', text)
        
        # Normalize paragraph breaks - convert 3+ newlines to double newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove lines that are just whitespace
        lines = []
        for line in text.split('\n'):
            cleaned_line = line.strip()
            if cleaned_line:  # Only keep non-empty lines
                lines.append(cleaned_line)
            else:
                # Preserve paragraph breaks
                if lines and lines[-1] != '':
                    lines.append('')
        
        # Rejoin and clean up
        text = '\n'.join(lines)
        
        # Remove any remaining excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()

    def _extract_from_pdf(self, file) -> str:
        """Extract text from PDF with better structure preservation."""
        try:
            pdf_reader = PdfReader(io.BytesIO(file.read()))
            text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        # Clean page text
                        page_text = page_text.strip()
                        # Add page separator for context
                        text_parts.append(f"Page {page_num + 1}:\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            raise DocumentProcessingError(f"PDF processing failed: {str(e)}")

    def _extract_from_docx(self, file) -> str:
        """Extract text from DOCX with better structure preservation."""
        try:
            doc = docx.Document(io.BytesIO(file.read()))
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(' | '.join(row_text))
                
                if table_text:
                    text_parts.append('\n'.join(table_text))
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            raise DocumentProcessingError(f"DOCX processing failed: {str(e)}")

    def _extract_from_txt(self, file) -> str:
        """Extract text from TXT file with encoding detection."""
        try:
            encodings = [self.file_config.ENCODING, 'utf-8', 'latin-1', 'cp1252']
            for encoding in encodings:
                try:
                    file.seek(0)
                    content = file.getvalue()
                    if isinstance(content, bytes):
                        text = content.decode(encoding)
                    else:
                        text = content
                    return text
                except UnicodeDecodeError:
                    continue
            raise DocumentProcessingError("Could not decode text file with any supported encoding")
        except Exception as e:
            raise DocumentProcessingError(f"TXT processing failed: {str(e)}")

    def create_chunks(self, text: str) -> List[str]:
        """Create semantic-aware chunks with enhanced filtering."""
        if not text or not text.strip():
            return []

        try:
            # Pre-process text for better chunking
            processed_text = self._preprocess_for_chunking(text)
            
            # Split into chunks
            raw_chunks = self.text_splitter.split_text(processed_text)
            
            # Enhanced chunk filtering and validation
            quality_chunks = self._filter_and_enhance_chunks(raw_chunks)
            
            # Limit chunks if necessary
            if len(quality_chunks) > self.ai_config.MAX_CHUNKS:
                quality_chunks = self._select_best_chunks(quality_chunks, self.ai_config.MAX_CHUNKS)

            logger.info(f"Created {len(quality_chunks)} high-quality chunks from {len(text)} characters")
            return quality_chunks

        except Exception as e:
            logger.error(f"Text chunking failed: {e}")
            raise DocumentProcessingError(f"Failed to create text chunks: {str(e)}")

    def _preprocess_for_chunking(self, text: str) -> str:
        """Preprocess text to improve chunking quality."""
        import re
        
        # Ensure sentences end with proper punctuation for better splitting
        text = re.sub(r'([a-zA-Z0-9])\s*\n\s*([A-Z])', r'\1. \2', text)
        
        # Fix common formatting issues
        text = re.sub(r'([.!?])\s*([A-Z])', r'\1 \2', text)
        
        return text

    def _filter_and_enhance_chunks(self, chunks: List[str]) -> List[str]:
        """Enhanced chunk filtering with quality scoring."""
        quality_chunks = []
        
        for i, chunk in enumerate(chunks):
            chunk = chunk.strip()
            
            # Skip chunks that are too short to be meaningful
            if len(chunk) < 30:
                continue
            
            # Skip chunks with no meaningful content
            if not self._has_meaningful_content(chunk):
                continue
            
            # Skip duplicate or near-duplicate chunks
            if self._is_duplicate_chunk(chunk, quality_chunks):
                continue
            
            # Add context-preserving prefix for better coherence
            enhanced_chunk = self._enhance_chunk_context(chunk, chunks, i)
            
            quality_chunks.append(enhanced_chunk)
        
        return quality_chunks

    def _has_meaningful_content(self, chunk: str) -> bool:
        """Check if chunk has meaningful content."""
        # Must have some alphanumeric content
        if not any(c.isalnum() for c in chunk):
            return False
        
        # Must have some substantial words
        words = chunk.split()
        meaningful_words = [w for w in words if len(w) > 2 and any(c.isalnum() for c in w)]
        
        if len(meaningful_words) < 3:
            return False
        
        # Should not be mostly punctuation or formatting
        alpha_ratio = sum(c.isalpha() for c in chunk) / len(chunk)
        if alpha_ratio < 0.3:
            return False
        
        return True

    def _is_duplicate_chunk(self, chunk: str, existing_chunks: List[str]) -> bool:
        """Check if chunk is too similar to existing chunks."""
        chunk_words = set(chunk.lower().split())
        
        for existing in existing_chunks[-3:]:  # Check last 3 chunks for efficiency
            existing_words = set(existing.lower().split())
            
            # Calculate similarity
            intersection = len(chunk_words.intersection(existing_words))
            union = len(chunk_words.union(existing_words))
            
            if union > 0:
                similarity = intersection / union
                if similarity > 0.8:  # 80% similarity threshold
                    return True
        
        return False

    def _enhance_chunk_context(self, chunk: str, all_chunks: List[str], index: int) -> str:
        """Add context to chunks for better understanding."""
        # For very short chunks, add context from previous chunk
        if len(chunk) < 100 and index > 0:
            prev_chunk = all_chunks[index - 1].strip()
            if prev_chunk and len(prev_chunk) > 50:
                # Add last sentence or phrase from previous chunk as context
                context_sentences = prev_chunk.split('. ')
                if len(context_sentences) > 1:
                    context = context_sentences[-1].strip()
                    if len(context) < 100:  # Don't add too much context
                        chunk = f"[Previous context: ...{context}] {chunk}"
        
        return chunk

    def _select_best_chunks(self, chunks: List[str], max_chunks: int) -> List[str]:
        """Select the best chunks when we have too many."""
        if len(chunks) <= max_chunks:
            return chunks
        
        # Score chunks based on various factors
        scored_chunks = []
        for i, chunk in enumerate(chunks):
            score = self._score_chunk_quality(chunk, i, len(chunks))
            scored_chunks.append((score, chunk))
        
        # Sort by score and take the top chunks
        scored_chunks.sort(key=lambda x: x[0], reverse=True)
        
        return [chunk for _, chunk in scored_chunks[:max_chunks]]

    def _score_chunk_quality(self, chunk: str, position: int, total_chunks: int) -> float:
        """Score chunk quality for selection."""
        score = 0.0
        
        # Length score (prefer medium-length chunks)
        length = len(chunk)
        if 200 <= length <= 600:
            score += 1.0
        elif 100 <= length < 200 or 600 < length <= 800:
            score += 0.7
        else:
            score += 0.3
        
        # Content diversity score
        words = chunk.split()
        unique_words = len(set(word.lower() for word in words))
        if words:
            diversity = unique_words / len(words)
            score += diversity
        
        # Position score (slightly prefer chunks from beginning and middle)
        relative_position = position / total_chunks
        if relative_position < 0.7:  # First 70% of document
            score += 0.3
        
        # Sentence completeness score
        if chunk.strip().endswith(('.', '!', '?')):
            score += 0.2
        
        return score

    def get_document_stats(self, text: str, chunks: List[str]) -> dict:
        """Get enhanced statistics about processed document."""
        return {
            'original_length': len(text),
            'total_chunks': len(chunks),
            'average_chunk_size': sum(len(chunk) for chunk in chunks) // len(chunks) if chunks else 0,
            'min_chunk_size': min(len(chunk) for chunk in chunks) if chunks else 0,
            'max_chunk_size': max(len(chunk) for chunk in chunks) if chunks else 0,
            'total_processed_length': sum(len(chunk) for chunk in chunks),
            'chunk_overlap_ratio': self.ai_config.CHUNK_OVERLAP / self.ai_config.CHUNK_SIZE if self.ai_config.CHUNK_SIZE > 0 else 0,
            'processing_efficiency': sum(len(chunk) for chunk in chunks) / len(text) if text else 0
        }

# Factory function
def create_document_processor() -> DocumentProcessor:
    """Create a new enhanced document processor instance"""
    return DocumentProcessor()

# Convenience functions for backward compatibility
def get_document_text(file):
    """Extract text from document file (backward compatibility)"""
    processor = create_document_processor()
    text, _ = processor.extract_text(file)
    return text

def get_text_chunks(text: str) -> List[str]:
    """Create text chunks from text (backward compatibility)"""
    processor = create_document_processor()
    return processor.create_chunks(text)